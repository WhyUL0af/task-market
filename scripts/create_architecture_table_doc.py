from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/task-market-architecture-table.docx")


def set_run(run, size=11, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph(paragraph, after=6):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run(run, size=26)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run(run, size=16)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=8)
    run = paragraph.add_run(text)
    set_run(run)


def add_code_block(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()


def fill_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run(run, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        fill_cell(header_cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            fill_cell(cells[index], value)

    widths = [1.45, 1.85, 3.2]
    for row in table.rows:
        for index, width in enumerate(widths[: len(headers)]):
            row.cells[index].width = Inches(width)

    doc.add_paragraph()
    return table


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    add_title(doc, "系統架構表")
    add_body(doc, "這份表格整理任務發布系統用到的主要技術，以及每個部分負責做什麼。")

    add_heading(doc, "主要架構")
    add_table(
        doc,
        ["層級", "使用技術", "主要作用"],
        [
            ["前端", "Next.js / React / TypeScript", "顯示網頁畫面，讓使用者操作任務、申請、提交成果。"],
            ["後端", "NestJS / TypeScript", "處理 API、任務流程、權限、審核和驗收。"],
            ["資料庫", "PostgreSQL", "儲存使用者、任務、申請、提交、留言、EXP 和徽章。"],
            ["ORM", "Prisma", "讓後端用 TypeScript 操作資料庫。"],
            ["登入驗證", "JWT", "確認使用者身份，以及判斷 Admin 或 Employee。"],
            ["Docker", "Docker Compose", "在本機啟動 PostgreSQL，讓環境比較固定。"],
        ],
    )

    add_heading(doc, "使用流程")
    add_table(
        doc,
        ["流程", "經過哪裡", "說明"],
        [
            ["打開網站", "瀏覽器 -> 前端", "使用者看到 Next.js 做出的網頁。"],
            ["登入", "前端 -> 後端 API", "帳號密碼送到 NestJS API 驗證。"],
            ["取得身份", "後端 -> JWT", "登入成功後產生 token。"],
            ["操作任務", "前端 -> API", "使用者申請、提交或審核任務。"],
            ["存取資料", "API -> Prisma -> PostgreSQL", "後端把資料存進資料庫或從資料庫讀出來。"],
            ["顯示結果", "API -> 前端", "前端把任務資料、狀態和結果顯示出來。"],
        ],
    )

    add_heading(doc, "簡單架構圖")
    add_code_block(
        doc,
        [
            "```mermaid",
            "flowchart TD",
            "    A[使用者瀏覽器] --> B[Next.js 前端]",
            "    B --> C[NestJS API]",
            "    C --> D[Prisma ORM]",
            "    D --> E[PostgreSQL 資料庫]",
            "",
            "    C --> F[JWT 驗證]",
            "    C --> G[任務流程]",
            "    C --> H[遊戲化功能]",
            "",
            "    G --> G1[任務建立]",
            "    G --> G2[任務申請]",
            "    G --> G3[申請審核]",
            "    G --> G4[成果提交]",
            "    G --> G5[驗收與結案]",
            "",
            "    H --> H1[EXP]",
            "    H --> H2[稱號]",
            "    H --> H3[徽章]",
            "    H --> H4[排行榜]",
            "```",
        ],
    )

    add_heading(doc, "一句話")
    add_body(
        doc,
        "這個系統是用 Next.js 做畫面、NestJS 做 API、PostgreSQL 存資料，Prisma 負責連資料庫，JWT 負責登入驗證。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
