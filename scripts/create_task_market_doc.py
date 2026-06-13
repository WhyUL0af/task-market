from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/task-market-system-overview.docx")


def set_run(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run(run, size=26, bold=False)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(20 if level == 1 else 18)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run(run, size=20 if level == 1 else 16, bold=False)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run(run)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run(run)


def add_steps(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run(run)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_document(doc)

    add_title(doc, "任務發布系統功能說明")
    add_body(
        doc,
        "這是一個讓團隊發布任務、申請任務、提交成果和驗收成果的系統。"
        "它也加了一些 EXP、稱號、徽章和排行榜，讓做任務比較有成就感。"
    )

    add_heading(doc, "1. 這個系統在做什麼")
    add_body(
        doc,
        "簡單來說，Admin 可以把任務放到系統上，員工可以自己申請想做的任務。"
        "Admin 看過申請後決定錄取誰，員工完成後提交成果，最後由 Admin 驗收和結案。"
    )

    add_heading(doc, "2. 系統角色")
    add_body(doc, "系統主要有兩種角色。")
    add_bullets(
        doc,
        [
            "Admin：可以新增任務、審核申請、錄取人員、驗收成果、結案任務，也可以管理使用者和技能標籤。",
            "Employee：可以看任務、申請任務、查看自己的任務、提交成果，也可以看自己的 EXP、稱號和徽章。",
        ],
    )

    add_heading(doc, "3. 任務可以設定什麼")
    add_body(doc, "每個任務可以設定一些基本資料。")
    add_bullets(
        doc,
        [
            "任務名稱",
            "任務說明",
            "預算",
            "EXP 獎勵",
            "難度",
            "截止時間",
            "需要哪些技能",
            "需要幾個人",
        ],
    )

    add_heading(doc, "4. 招募需求")
    add_body(
        doc,
        "一個任務可以拆成多個需求。像是一個任務要兩個人，可以設定一個人需要前端技能，"
        "另一個人需要後端技能。"
    )
    add_bullets(
        doc,
        [
            "前端需求：1 人，預算 50%，EXP 50%。",
            "後端需求：1 人，預算 50%，EXP 50%。",
            "全部需求的預算比例加起來要等於 100%。",
            "全部需求的 EXP 比例加起來也要等於 100%。",
        ],
    )

    add_heading(doc, "5. 任務流程")
    add_steps(
        doc,
        [
            "Admin 發布任務。",
            "員工選擇一個需求並送出申請。",
            "Admin 審核申請。",
            "Admin 錄取適合的人。",
            "員工開始做任務。",
            "員工提交成果。",
            "Admin 驗收成果。",
            "如果成果不行，Admin 可以退回修改。",
            "所有錄取人員都驗收通過後，Admin 結案。",
            "系統統一發放 EXP。",
        ],
    )

    add_heading(doc, "6. 任務狀態")
    add_bullets(
        doc,
        [
            "Draft：草稿",
            "Open：招募中",
            "Applied：已有人申請",
            "In Progress：進行中",
            "Review：驗收中",
            "Done：已完成",
            "Cancelled：已取消",
        ],
    )

    add_heading(doc, "7. 遊戲化功能")
    add_body(doc, "系統有一些簡單的遊戲化功能。")
    add_bullets(
        doc,
        [
            "EXP：完成任務後會獲得經驗值。",
            "稱號：EXP 累積後會提升稱號。",
            "徽章：達成特定條件後可以解鎖。",
            "每週挑戰：每週有一些小目標。",
            "排行榜：可以看到 EXP、完成任務、準時率和協作排名。",
        ],
    )

    add_heading(doc, "8. 稱號")
    add_body(doc, "稱號最高到傳奇。")
    add_bullets(
        doc,
        ["新秀", "能手", "好手", "達人", "專家", "菁英", "首席", "傳奇"],
    )

    add_heading(doc, "9. 個人頁面")
    add_body(doc, "使用者可以在個人頁面看到自己的資料和任務表現。")
    add_bullets(
        doc,
        [
            "基本資料",
            "技能標籤",
            "通知設定",
            "任務統計",
            "EXP",
            "稱號",
            "徽章",
            "最近 EXP 紀錄",
        ],
    )

    add_heading(doc, "10. 管理功能")
    add_body(doc, "Admin 可以管理系統裡的人和標籤。")
    add_bullets(
        doc,
        [
            "新增使用者",
            "編輯使用者資料",
            "重設密碼",
            "修改 EXP 和稱號等級",
            "設定使用者技能標籤",
            "刪除帳號",
            "新增、編輯、刪除技能標籤",
        ],
    )

    add_heading(doc, "一句話總結")
    add_body(
        doc,
        "這個系統就是讓團隊可以把任務放出來，讓員工自己申請，Admin 再挑人、驗收成果，"
        "最後用 EXP、稱號和徽章增加一點成就感。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
