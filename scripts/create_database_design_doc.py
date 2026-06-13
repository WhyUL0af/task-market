from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/task-market-database-design.docx")


def set_run(run, size=11, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run(run, size=26)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run(run, size=16)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run(run)


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run(run)


def fill_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run(run, bold=bold)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for index, header in enumerate(headers):
      fill_cell(table.rows[0].cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            fill_cell(cells[index], value)

    if widths:
        for table_row in table.rows:
            for index, width in enumerate(widths[: len(headers)]):
                table_row.cells[index].width = Inches(width)

    doc.add_paragraph()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    add_title(doc, "資料庫設計")
    add_body(doc, "這份文件整理任務發布系統的主要資料表、欄位和資料表之間的關係。")

    add_heading(doc, "主要資料表")
    add_table(
        doc,
        ["資料表", "用途"],
        [
            ["User", "儲存使用者資料。"],
            ["Task", "儲存任務資料。"],
            ["TaskRequirement", "儲存任務的招募需求。"],
            ["TaskRequirementSkill", "儲存每個招募需求需要的技能。"],
            ["TaskApplication", "儲存員工的任務申請。"],
            ["TaskSubmission", "儲存員工提交的成果。"],
            ["Comment", "儲存任務留言。"],
            ["ProfileTag", "儲存技能標籤。"],
            ["UserProfileTag", "儲存使用者擁有的技能。"],
            ["Badge / UserBadge", "儲存徽章，以及使用者已獲得的徽章。"],
            ["Title / UserTitle", "儲存稱號，以及使用者已獲得的稱號。"],
            ["ExpTransaction", "儲存 EXP 增加紀錄。"],
            ["WeeklyChallenge", "儲存每週挑戰。"],
            ["UserChallengeProgress", "儲存使用者每週挑戰進度。"],
        ],
        widths=[2.1, 4.4],
    )

    add_heading(doc, "User 使用者表")
    add_table(
        doc,
        ["欄位", "說明"],
        [
            ["id", "使用者 ID。"],
            ["email", "登入帳號。"],
            ["name", "使用者名稱。"],
            ["passwordHash", "加密後的密碼。"],
            ["role", "使用者角色，Admin 或 Employee。"],
            ["bio", "基本資料。"],
            ["xp", "累積 EXP。"],
            ["level", "稱號等級。"],
            ["notificationSettings", "通知設定。"],
            ["createdAt / updatedAt", "建立時間與更新時間。"],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "Task 任務表")
    add_table(
        doc,
        ["欄位", "說明"],
        [
            ["id", "任務 ID。"],
            ["title", "任務名稱。"],
            ["description", "任務說明。"],
            ["reward", "任務預算。"],
            ["xpReward", "任務 EXP 獎勵。"],
            ["difficulty", "任務難度。"],
            ["dueAt", "截止時間。"],
            ["status", "任務狀態。"],
            ["creatorId", "建立任務的 Admin。"],
            ["assigneeId", "舊版主要指派人員欄位。"],
            ["createdAt / updatedAt", "建立時間與更新時間。"],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "TaskRequirement 招募需求表")
    add_body(doc, "一個任務可以有多個招募需求，例如前端需求 1 人、後端需求 1 人。")
    add_table(
        doc,
        ["欄位", "說明"],
        [
            ["id", "招募需求 ID。"],
            ["taskId", "所屬任務。"],
            ["name", "需求名稱。"],
            ["headcount", "需求人數。"],
            ["budgetPercent", "預算比例。"],
            ["xpPercent", "EXP 比例。"],
            ["createdAt", "建立時間。"],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "TaskApplication 任務申請表")
    add_table(
        doc,
        ["欄位", "說明"],
        [
            ["id", "申請 ID。"],
            ["taskId", "申請的任務。"],
            ["applicantId", "申請人。"],
            ["requirementId", "申請的招募需求。"],
            ["message", "申請訊息。"],
            ["status", "申請狀態。"],
            ["skillMatchScore", "技能匹配分數。"],
            ["assignedBudget", "分配到的預算。"],
            ["assignedXp", "分配到的 EXP。"],
            ["completedAt", "完成時間。"],
            ["onTime", "是否準時。"],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "TaskSubmission 任務提交表")
    add_table(
        doc,
        ["欄位", "說明"],
        [
            ["id", "提交 ID。"],
            ["taskId", "所屬任務。"],
            ["employeeId", "提交者。"],
            ["content", "成果內容。"],
            ["status", "提交狀態，包含待驗收、已驗收、退回修改。"],
            ["createdAt / updatedAt", "建立時間與更新時間。"],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "技能與遊戲化資料表")
    add_table(
        doc,
        ["資料表", "說明"],
        [
            ["ProfileTag", "儲存技能標籤。"],
            ["UserProfileTag", "記錄使用者有哪些技能。"],
            ["TaskRequirementSkill", "記錄每個招募需求需要哪些技能。"],
            ["Badge", "儲存徽章。"],
            ["UserBadge", "記錄使用者已獲得哪些徽章。"],
            ["Title", "儲存稱號。"],
            ["UserTitle", "記錄使用者已獲得哪些稱號。"],
            ["ExpTransaction", "記錄 EXP 增加紀錄。"],
            ["WeeklyChallenge", "儲存每週挑戰。"],
            ["UserChallengeProgress", "記錄使用者每週挑戰進度。"],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "資料表關係")
    add_bullets(
        doc,
        [
            "一個 Admin 可以建立多個任務。",
            "一個任務可以有多個招募需求。",
            "一個招募需求可以需要多個技能。",
            "一個使用者可以有多個技能，一個技能也可以被多個使用者擁有。",
            "一個任務可以有多個申請。",
            "一個使用者可以申請多個任務。",
            "一個任務可以有多個提交紀錄。",
            "一個任務可以有多個留言。",
            "一個使用者可以有多筆 EXP 紀錄、徽章、稱號和挑戰進度。",
        ],
    )

    add_heading(doc, "一句話總結")
    add_body(
        doc,
        "這個資料庫主要以 User 和 Task 為核心。Task 延伸出招募需求、申請、提交和留言；"
        "User 延伸出技能、EXP、徽章、稱號和每週挑戰。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
